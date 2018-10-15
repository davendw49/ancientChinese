# -*- coding: utf-8 -*-
import codecs
import MySQLdb
import parse_html
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

reload(sys);
sys.setdefaultencoding("utf8")



def query(chinese, table_name):
	# 连接
	db = MySQLdb.connect("localhost", "root", "kellydc", "ancientChinese", charset='utf8' )
	# 获取操作游标
	cursor=db.cursor()
	# 使用execute方法执行SQL语句

	result_set = []

	n = cursor.execute("select * from %s where chinese='%s'" % (table_name, chinese))
	for info in cursor.fetchall():
		result_set.append(info)

	return result_set


print len(parse_html.chinese_set)
ids = list(set(parse_html.chinese_set))
print len(ids)

def html_output(html_file):
	f = open(html_file, 'w+')
	length = len(ids)
	pace = int(length/100)+1
	nowpace = 0
	nowpercant = 0
	for charI in ids:
		nowpace+=1
		print >> f,"<p>"
		print >> f,charI
		print >> f,"</p>"
		result = query(charI, "sample")
		print >> f,"<p>"
		for rs in result:
			print >> f,rs[3],rs[2],","
		print >> f,"</p>"
		if nowpace==pace:
			nowpercant+=1
			print str(nowpercant)+"%"
			nowpace = 0

	f.close()
	print "done"

def word_output(word_file):
	document = Document()
	document.styles['Normal'].font.name = u'黑体'
	document.add_heading(word_file, 0)
	length = len(ids)
	pace = int(length/100)+1
	nowpace = 0
	nowpercant = 0
	for charI in ids:
		nowpace+=1
		p1 = document.add_paragraph(charI)
		result = query(charI, "sample")
		p = document.add_paragraph()
		for rs in result:
			run = p.add_run(str(rs[2]))
			run.add_picture(str(rs[3])[7:])
			run.add_text(';')
		if nowpace==pace:
			nowpercant+=1
			print str(nowpercant)+"%"
			nowpace = 0
	document.save(word_file)  #可以设置其他路径


if sys.argv[1]=="docx":
	word_output(sys.argv[2])
elif sys.argv[1]=="html":
	html_output(sys.argv[2])




